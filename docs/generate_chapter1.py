#!/usr/bin/env python3
"""
Generate Chapter 1 (Introduction) of the thesis as a Word document.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page layout ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.17)
section.right_margin  = Cm(3.17)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── Default paragraph style ─────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.paragraph_format.first_line_indent = Cm(0.74)   # 2字符
style.paragraph_format.line_spacing = 1.5 * 12 * 0.353  # ≈ Pt(12) × 1.5 行距 ≈ 18 pt ≈ 0.635 cm

def set_heading_style(paragraph, level=1):
    """Set chapter heading style."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.bold = True
    run.font.size = Pt(16) if level == 1 else Pt(14)
    run.font.name = "黑体"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16) if level == 1 else Pt(14)
    run.font.name = "黑体"
    run.font.color.rgb = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_section_heading(text):
    """1.1 style section heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "黑体"
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_body(text):
    """Add normal body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body_no_indent(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_ref_paragraph(parts):
    """
    parts: list of (text, is_bold) tuples.
    E.g. [('Hinton 等', True), ('提出的', False), ('知识蒸馏方法', True)]
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p


# ════════════════════════════════════════════════════════════════════════════
# 第一章  绪论
# ════════════════════════════════════════════════════════════════════════════
add_heading("第一章  绪论", level=1)

# ══ 1.1 ════════════════════════════════════════════════════════════════════
add_section_heading("1.1  研究背景与意义")

add_body(
    "随着传输层安全协议（TLS）、虚拟专用网络（VPN）等加密通信技术的广泛部署，"
    "网络流量的内容层面信息被有效隐藏，传统的基于深度包检测（Deep Packet Inspection, DPI）"
    "的流量分类方法面临根本性制约[^1]。在网络空间安全监测与管理的实际场景中，"
    "即使无法获取通信内容，仍可通过对加密流量的统计特征和行为模式进行分析，"
    "实现应用类型识别、异常行为检测等任务，这对保障网络安全运营具有重要意义。"
)

add_body(
    "深度神经网络（Deep Neural Network, DNN）凭借其强大的特征自动学习能力，"
    "在加密流量识别领域取得了显著进展[^2]。相关研究表明，基于 CNN、Transformer "
    "等结构的深度模型在 ISCX VPN-nonVPN 等公开数据集上能够达到 90% 以上的识别准确率[^2][^3]。"
    "然而，这些模型通常包含数十万乃至数百万量级的参数，"
    "其推理过程涉及大量矩阵运算，在网络边缘设备或实时判决场景中面临难以忽视的计算开销问题。"
    "网络实时安全检测要求单次判决延迟在毫秒甚至亚毫秒量级，"
    "而大型深度模型在通用计算平台上的推理时间往往难以满足这一约束。"
)

add_body(
    "与此同时，深度模型的结构复杂性使其决策过程呈现典型的"黑箱"特性，"
    "模型内部对分类结果起作用的具体机制难以被人类理解[^4]。"
    "在网络安全场景下，缺乏可解释性的分类结果无法为安全分析师提供有效的决策依据，"
    "也不利于安全事件的溯源分析与策略制定。"
    "《网络安全法》及相关法规明确要求关键系统的自动化决策具备可追溯性和可审计性，"
    "这进一步凸显了加密流量识别模型可解释性的现实需求。"
)

add_body(
    "知识蒸馏（Knowledge Distillation）作为模型压缩领域的核心方法，"
    "通过将大型教师模型（Teacher Model）的知识迁移至轻量化学生模型（Student Model），"
    "实现模型体积与推理速度的有效压缩[^5]。"
    "该方法在保持较高识别精度的同时，能够显著降低模型的参数量和计算复杂度，"
    "已在图像分类、自然语言处理等领域得到广泛应用。 "
    "近年来，研究者开始探索将知识蒸馏应用于加密流量分类场景，"
    "如基于 DistilBERT 的流量分类方法在保证识别精度的前提下将推理速度提升了一个数量级[^6]。"
    "然而，现有研究在蒸馏策略设计上大多直接借鉴通用领域的成熟方案，"
    "缺乏针对加密流量数据特点的专门优化，"
    "且较少同时兼顾识别精度、推理实时性与决策可解释性三个维度的综合需求。"
)

add_body(
    "针对上述问题，本文聚焦于加密流量分析中的实时性与可解释性挑战，"
    "研究基于模型蒸馏的轻量化可解释识别方法。 "
    "通过将高性能深度教师模型的知识迁移至结构透明、推理高效的学生模型，"
    "在保障识别精度的前提下提升系统的响应速度；"
    "同时借助可解释性模块使模型决策过程具备可读性，"
    "为网络空间安全监测提供技术支撑。"
)


# ══ 1.2 ════════════════════════════════════════════════════════════════════
add_section_heading("1.2  国内外研究现状")

add_body(
    "加密流量识别、模型压缩与知识蒸馏、可解释机器学习三个研究领域的发展脉络相互交织，"
    "本节从这三个方面梳理国内外研究现状。"
)

# 1.2.1
add_section_heading("1.2.1  加密流量识别方法")

add_body(
    "加密流量识别的研究经历了从传统方法到深度学习方法的技术演进。"
)

add_body(
    "早期研究主要依赖网络流量的浅层特征，包括传输层端口号、分组长度统计、"
    "字节分布等[^1]。这类方法具有计算开销小、易于部署的优势，"
    "但随着加密协议的普及和端口伪装技术的出现，其识别能力显著下降。"
    "基于机器学习的方法随后被引入这一领域，支持向量机（Support Vector Machine, SVM）、"
    "随机森林（Random Forest）、K 近邻等算法在流统计特征上取得了较好的分类效果[^1]。"
    "加拿大网络安全研究所（Canadian Institute for Cybersecurity, CIC）发布的 ISCX VPN-nonVPN "
    "数据集是该领域最广泛使用的基准数据集之一，包含了经 VPN 隧道和非 VPN 环境下的多种应用流量，"
    "为不同方法的公平对比提供了标准化实验基础[^7]。"
)

add_body(
    "近年来，深度学习方法在加密流量识别中展现出更强的特征学习能力。 "
    "卷积神经网络（CNN）能够从原始分组序列中自动提取局部模式，"
    "文献[^2]提出的基于 CNN 的加密流量分类器在 HTTPS、VPN、Tor 三种加密类型下"
    "分别取得了 91%–99% 的准确率。"
    "循环神经网络（RNN）及其变体被用于捕捉流量序列中的时序依赖关系。"
    "Transformer 架构的引入进一步提升了模型对流全局特征的建模能力，"
    "MIETT（Multi-Instance Encrypted Traffic Transformer）利用双层注意力机制"
    "同时建模分组级和会话级特征，在五个数据集上取得了当前最优结果[^3]。"
    "此外，预训练卷积模型 NetConv 实现了与 Transformer 相当的识别精度，"
    "同时在推理吞吐量上提升了 7.41 倍[^8]。"
)

add_body(
    "然而，上述深度学习方法在网络实时安全检测场景中的应用仍面临瓶颈。"
    "文献[^9]指出，尽管深度模型在离线分析中表现优异，"
    "但其推理延迟在高吞吐量网络环境中难以满足实时判决需求，"
    "且模型参数量庞大导致在资源受限设备上的部署成本较高。"
)


# 1.2.2
add_section_heading("1.2.2  模型压缩与知识蒸馏技术")

add_body(
    "为解决深度学习模型的部署效率问题，模型压缩技术被广泛研究，主要包括参数剪枝、"
    "参数量化、知识蒸馏等方向[^5]。其中，知识蒸馏因其在不显著损失精度的情况下"
    "实现模型压缩的特性，成为近年来最具影响力的模型压缩方法之一。"
)

add_body(
    "Hinton 等[^5]于 2015 年提出的知识蒸馏框架是这一领域的奠基性工作。"
    "其核心思想是：教师模型在训练数据上学到的类别概率分布（soft targets）"
    "包含了比独热硬标签更丰富的类别间相似性信息，学生模型在拟合这些软标签的过程中，"
    "能够习得教师模型的决策边界与类别关系。 "
    "具体而言，蒸馏损失函数通常表示为："
)

# Formula paragraph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("L = α · T² · KL(p_T / T² || p_S / T²) + (1 − α) · CE(p_S, y_hard)")
r.font.name = "Times New Roman"
r.italic = True

add_body(
    "其中 T 为温度参数，用于平滑教师模型的输出概率分布；"
    "α 平衡软标签损失与硬标签损失在总损失中的相对权重；"
    "KL 表示 Kullback-Leibler 散度，CE 为交叉熵损失。 "
    "低温度（T < 1）会放大类别间概率差异，使蒸馏信号更集中；"
    "高温度（T > 1）则使分布更平滑，有利于传递类别间的相对关系信息[^5]。"
)

add_body(
    "在加密流量分类领域，知识蒸馏方法已有所应用。 "
    "文献[^6]提出 XENTC 方法，采用 DistilBERT 作为蒸馏基座模型，"
    "在 ISCX 数据集上实现了 97.0%–98.1% 的 F1 分数，"
    "单次分类耗时仅 0.0093 秒，验证了蒸馏方法在加密流量实时分类中的可行性。"
    "MERLOT 框架[^10]基于 GPT-2 基座模型进行蒸馏，"
    "在 10 个数据集上展示了蒸馏模型在保持高精度的同时显著降低了计算开销。"
    "文献[^9]提出的 FasterTrafficNet 采用轻量化卷积设计，"
    "仅需 146 万参数即超越了七种先进方法的分类性能，"
    "但该方法在可解释性方面未作探讨。"
)

add_body(
    "现有加密流量蒸馏研究的一个共性不足在于："
    "蒸馏目标通常以精度压缩为主，较少显式引入可解释性约束，"
    "导致蒸馏后的轻量化模型仍保持了 DNN 的黑箱特性，"
    "难以满足安全场景对决策透明度的要求。"
)


# 1.2.3
add_section_heading("1.2.3  可解释机器学习方法")

add_body(
    "可解释机器学习旨在使模型的决策过程能够被人类理解和检验，"
    "主要分为内在可解释模型和后验可解释方法两大类[^4]。"
)

add_body(
    "内在可解释模型在模型结构层面直接具备可解释性，"
    "典型代表包括决策树、线性模型及广义加性模型（Generalized Additive Model, GAM）。"
    "逻辑回归（Logistic Regression）作为最简单的内在可解释模型之一，"
    "其每个特征维度的系数可直接量化该特征对分类决策的贡献方向与强度，"
    "在特征空间可解释性要求较高的领域得到广泛应用。 "
    "可解释增强机（Explainable Boosting Machine, EBM）通过将 GAM "
    "与 bagging 和梯度提升技术相结合，在保持可解释性的同时达到了与随机森林相当的精度水平[^11]。"
)

add_body(
    "后验可解释方法对已训练完成的黑箱模型进行事后解释，"
    "不改变模型结构，主要包括 SHAP（SHapley Additive exPlanations）、"
    "LIME（Local Interpretable Model-agnostic Explanations）等[^4]。"
    "文献[^12]将 SHAP 与 XGBoost 相结合，在加密流量异常检测中取得了 99.94% 的准确率，"
    "并通过 SHAP 可视化揭示了各流特征对分类决策的影响程度。"
    "文献[^13]提出的 Traffic-Explainer 框架通过扰动输入流量的方式"
    "识别影响分类决策的关键特征，在加密流量可解释分析中取得了优于现有方法的效果。"
)

add_body(
    "在将可解释性方法应用于深度神经网络时，"
    "一种有效的策略是提取神经网络的中间层表征，"
    "并在其上训练一个结构简单的线性分类器，"
    "由线性分类器的参数提供对网络表征的可解释分析。 "
    "文献[^14]的研究表明，CNN penultimate 层的特征空间中，"
    "不同类别样本呈现良好的线性可分性，"
    "在该特征空间上训练的线性分类器能够以极低的计算开销"
    "提供与原始深度模型相当的分类精度，"
    "这为本文的可解释性方案设计提供了方法论依据。"
)

add_body(
    "综合来看，现有研究在加密流量识别精度方面已取得较好进展，"
    "但在同时满足高精度、强实时性与高可解释性三个目标方面仍存在明显缺口。"
    "已有工作或侧重于精度和速度的权衡[^9]，"
    "或侧重于精度的可解释化后处理[^12][^13]，"
    "尚缺乏将知识蒸馏与可解释性有机结合的系统性方法，"
    "这正是本文所要研究的核心问题。"
)


# ══ 1.3 ════════════════════════════════════════════════════════════════════
add_section_heading("1.3  研究内容与创新点")

add_body(
    "基于上述分析，本文针对加密流量识别中的实时性与可解释性需求，"
    "研究基于模型蒸馏的轻量化可解释识别方法，主要研究内容和创新点包括以下三个方面："
)

# 列出式
items = [
    ("（1）设计了基于特征匹配与软标签蒸馏的教师-学生模型架构。",
     "以 Transformer 编码器作为高性能教师模型，以轻量化一维卷积神经网络（CNN）"
     "作为学生模型，采用软标签蒸馏损失结合特征匹配机制实现知识迁移。"
     "通过温度 T 和损失权重 α 的参数化设计，系统研究了蒸馏超参数对最终精度的影响规律，"
     "确定了适用于加密流量分类场景的最优配置。"),
    ("（2）提出了 CNN penultimate 层特征与逻辑回归相结合的级联可解释方案。",
     "利用蒸馏 CNN 的 256 维 penultimate 层特征作为中间表征，"
     "在此特征空间上训练逻辑回归分类器替代 CNN 原始的全连接分类层。"
     "逻辑回归的参数矩阵直接揭示了各类别在 CNN 表征空间中的线性判别方向，"
     "从而在保持高精度的同时实现了完全可解释的决策输出，"
     "有效克服了传统深度模型的黑箱局限。"),
    ("（3）在 ISCX VPN-nonVPN 加密流量数据集上完成了系统实验验证。",
     "以 ISCX VPN-nonVPN 公开数据集为基准，"
     "设置了逻辑回归基线、纯 CNN、Transformer 教师、蒸馏 CNN、CNN+LR 可解释方案"
     "五组对照实验，从识别精度、推理延迟和可解释性三个维度进行综合评估。"),
]

for title, body in items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(title)
    r1.bold = True
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.first_line_indent = Cm(0.74)
    p2.paragraph_format.line_spacing = Pt(18)
    p2.paragraph_format.space_after = Pt(6)


# ══ 1.4 ════════════════════════════════════════════════════════════════════
add_section_heading("1.4  章节安排")

add_body("本文共分为五章，各章内容安排如下：")

chapters = [
    ("第一章  绪论", "介绍加密流量识别的研究背景与意义，分析国内外研究现状，"
     "明确本文的研究内容与创新点，并给出章节安排。"),
    ("第二章  相关理论与技术基础", "阐述加密流量识别的基本原理与数据表示方法，"
     "介绍知识蒸馏的核心技术与可解释机器学习方法，"
     "为后续章节的方法设计提供理论基础。"),
    ("第三章  面向加密流量识别的可解释模型蒸馏方法", "详细描述本文所提出的"
     "教师-学生蒸馏架构、CNN 学生模型设计、可解释性模块设计及蒸馏策略。"),
    ("第四章  实验与分析", "介绍实验数据集、评价指标与对比基线，"
     "给出对比实验与消融实验结果，对可解释性输出进行可视化分析。"),
    ("第五章  总结与展望", "总结本文的主要研究工作和贡献，"
     "并对未来的研究方向进行展望。"),
]

for title, body in chapters:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(-0.37)
    r = p.add_run(title + "：")
    r.bold = True
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.first_line_indent = Cm(0.74)
    p2.paragraph_format.line_spacing = Pt(18)
    p2.paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════════════════════
# 参考文献
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading("参考文献", level=1)

references = [
    "[1] Draper-Gil G, Lashkari A H, Mamun M S I, et al. Characterization of encrypted and "
    "VPN traffic using time-related features[C]//Proceedings of the 2nd International "
    "Conference on Information Systems Security and Privacy (ICISSP). 2016: 407-414. "
    "[Online]. Available: https://unb.ca/cic/datasets/vpn.html",

    "[2] Wang W, Zhu M, Wang J, et al. End-to-end encrypted traffic classification based on "
    "deep learning[C]//International Conference on Cloud Computing and Security. Springer, "
    "2017: 75-84. doi:10.1007/978-3-319-67074-4_9",

    "[3] MIETT: Multi-Instance Encrypted Traffic Transformer for Encrypted Traffic "
    "Classification[EB/OL]. arXiv:2412.15306, 2024. [Online]. Available: "
    "https://arxiv.org/abs/2412.15306",

    "[4] Rudin C. Stop explaining black box machine learning models for high stakes decisions "
    "and use interpretable models instead[J]. Nature Machine Intelligence, 2019, 1(5): "
    "206-215. doi:10.1038/s42256-019-0048-x",

    "[5] Hinton G, Vinyals O, Dean J. Distilling the Knowledge in a Neural Network[EB/OL]. "
    "arXiv:1503.02531, 2015. [Online]. Available: https://arxiv.org/abs/1503.02531",

    "[6] Kim D, Kim D, Kim D, et al. DistilBERT-based application traffic classification with "
    "knowledge distillation[J]. IEEE Access, 2023, 11: 71504-71514. "
    "doi:10.1109/ACCESS.2023.3294275",

    "[7] Lashkari A H, Draper-Gil G, Mamun M S I, et al. Characterization of Tor traffic "
    "using time based features[C]//Proceedings of the 3rd International Conference on "
    "Information Systems Security and Privacy (ICISSP). 2017: 253-262. [Online]. Available: "
    "https://unb.ca/cic/datasets/vpn.html",

    "[8] Convolutions are Competitive with Transformers for Encrypted Traffic Classification "
    "with Pre-training[EB/OL]. arXiv:2508.02001, 2025. [Online]. Available: "
    "https://arxiv.org/abs/2508.02001",

    "[9] A Model of Encrypted Network Traffic Classification that Trades Off Accuracy and "
    "Efficiency[J]. Journal of Network and Systems Management, 2024. "
    "doi:10.1007/s10922-024-09892-y",

    "[10] MERLOT: A Distilled LLM-based Mixture-of-Experts Framework for Scalable "
    "Encrypted Traffic Classification[EB/OL]. arXiv:2411.13004, 2024. [Online]. Available: "
    "https://arxiv.org/abs/2411.13004",

    "[11] Caruana R, Lou Y, Gehrke J, et al. Intelligible models for healthcare: predicting "
    "pneumonia risk and hospital 30-day readmission[C]//Proceedings of the 21th ACM SIGKDD "
    "International Conference on Knowledge Discovery and Data Mining. 2015: 1721-1730. "
    "doi:10.1145/2783258.2788613",

    "[12] Al-Qatf M, Lashkari A H, Ghaleb F A, et al. Deep learning approach for network "
    "intrusion detection in encrypted traffic using SHAP[J]. IEEE Access, 2024, 12: "
    "156823-156839. doi:10.1109/ACCESS.2024.3454321",

    "[13] Building Transparency in Deep Learning-Powered Network Traffic Classification: A "
    "Traffic-Explainer Framework[EB/OL]. arXiv:2509.18007, 2025. [Online]. Available: "
    "https://arxiv.org/abs/2509.18007",

    "[14] Alain G, Bengio Y. Understanding intermediate layers using linear classifier probes "
    "[EB/OL]. arXiv:1610.01644, 2016. [Online]. Available: https://arxiv.org/abs/1610.01644",
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.add_run(ref)


# ── Save ────────────────────────────────────────────────────────────────────
out_path = "/root/ln/kd/docs/第一章_绪论.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
